import streamlit as st
from docx import Document
from io import BytesIO
from docx.shared import Pt
from datetime import datetime

# Diccionario para asociar idioma con plantilla
PLANTILLAS = {
    "Español": "Carta Tipo - Incorporaciones.docx",
    "Portugués": "Carta Tipo - IncorporacionesPOR.docx",
    "Inglés": "Carta Tipo - IncorporacionesENG.docx",
}

# Diccionario de traducción de meses
MESES_TRADUCIDOS = {
    "Español": {
        "01": "Enero", "02": "Febrero", "03": "Marzo", "04": "Abril", "05": "Mayo", "06": "Junio",
        "07": "Julio", "08": "Agosto", "09": "Septiembre", "10": "Octubre", "11": "Noviembre", "12": "Diciembre"
    },
    "Portugués": {
        "01": "Janeiro", "02": "Fevereiro", "03": "Março", "04": "Abril", "05": "Maio", "06": "Junho",
        "07": "Julho", "08": "Agosto", "09": "Setembro", "10": "Outubro", "11": "Novembro", "12": "Dezembro"
    },
    "Inglés": {
        "01": "January", "02": "February", "03": "March", "04": "April", "05": "May", "06": "June",
        "07": "July", "08": "August", "09": "September", "10": "October", "11": "November", "12": "December"
    }
}

# Función para reemplazar solo los textos variables sin alterar el formato general
def reemplazar_campos(template_path, reemplazos):
    doc = Document(template_path)
    
    def reemplazar_texto_en_runs(runs, key, value):
        texto_completo = "".join(run.text for run in runs)
        if key in texto_completo:
            texto_modificado = texto_completo.replace(key, value)
            runs[0].text = texto_modificado
            for i in range(1, len(runs)):
                runs[i].text = ""
    
    for para in doc.paragraphs:
        for key, value in reemplazos.items():
            if key in para.text:
                reemplazar_texto_en_runs(para.runs, key, value)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in reemplazos.items():
                    if key in cell.text:
                        for para in cell.paragraphs:
                            reemplazar_texto_en_runs(para.runs, key, value)
    
    return doc

st.title("Generador de Carta de Incorporaciones")

# 🔧 Ocultar la barra superior y el menú de Streamlit
hide_streamlit_style = """
    <style>
        #MainMenu {visibility: hidden;}
        header {visibility: hidden;}
        footer {visibility: hidden;}
    </style>
"""
st.markdown(hide_streamlit_style, unsafe_allow_html=True)

idioma = st.selectbox("Seleccione el idioma", list(PLANTILLAS.keys()))

# Campos del formulario
nombre = st.text_input("Inserte Nombre")
localizador = st.text_input("Inserte Localizador")
fecha_input = st.text_input("Inserte Fecha (DD/MM/YYYY)")
ciudad = st.text_input("Inserte Ciudad")
trayecto = st.text_input("Inserte Trayecto")
hora_presentacion = st.text_input("Inserte Hora de Presentación")
hora_salida = st.text_input("Inserte Hora de Salida")
punto_encuentro = st.text_input("Inserte Punto de Encuentro")
direccion = st.text_area("Inserte Dirección")

# Validación de fecha y obtención del día y mes en texto
try:
    fecha_obj = datetime.strptime(fecha_input, "%d/%m/%Y")
    dia_semana = fecha_obj.strftime("%A")  # Día en inglés
    dia_num = fecha_obj.strftime("%d")
    mes_num = fecha_obj.strftime("%m")
    anio = fecha_obj.strftime("%Y")
    
    dias_traducidos = {
        "Español": {
            "Monday": "Lunes", "Tuesday": "Martes", "Wednesday": "Miércoles",
            "Thursday": "Jueves", "Friday": "Viernes", "Saturday": "Sábado", "Sunday": "Domingo"
        },
        "Portugués": {
            "Monday": "Segunda-feira", "Tuesday": "Terça-feira", "Wednesday": "Quarta-feira",
            "Thursday": "Quinta-feira", "Friday": "Sexta-feira", "Saturday": "Sábado", "Sunday": "Domingo"
        },
        "Inglés": {
            "Monday": "Monday", "Tuesday": "Tuesday", "Wednesday": "Wednesday",
            "Thursday": "Thursday", "Friday": "Friday", "Saturday": "Saturday", "Sunday": "Sunday"
        }
    }
    
    dia_traducido = dias_traducidos[idioma][dia_semana]
    mes_traducido = MESES_TRADUCIDOS[idioma][mes_num]
    fecha_formateada = f"{dia_num} - {mes_traducido} - {anio}"
    fecha_valida = True
except ValueError:
    st.error("Formato de fecha inválido. Use el formato DD/MM/YYYY.")
    fecha_valida = False

# Reemplazos si la fecha es válida
if fecha_valida and st.button("Generar Documento"):
    reemplazos = {
        "(INSERTENOMBRE)": nombre,
        "(LOCALIZADOR)": localizador,
        "(INSERTEFECHA)": fecha_formateada,
        "(DIA)": dia_traducido,
        "(CIUDAD)": ciudad,
        "(INSERTETRAYECTO)": trayecto,
        "(HORAPRESENTACION)": hora_presentacion,
        "(HORASALIDA)": hora_salida,
        "(PUNTODEENCUENTRO)": punto_encuentro,
        "(INSERTEDIRECCION)": direccion
    }

    plantilla = PLANTILLAS[idioma]
    doc = reemplazar_campos(plantilla, reemplazos)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.download_button(
        label="Descargar Documento",
        data=buffer,
        file_name=f"{localizador}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
